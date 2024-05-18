import tkinter
import base64
import tempfile
import os
import threading
import time
import calendar
import datetime
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from calendar_widget import Calendar


# Create a temporary image file
icon = "AAABAAEAICAAAAEAIACoEAAAFgAAACgAAAAgAAAAQAAAAAEAIAAAAAAAgBAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGJiYg12dnYrdnZ2K3Z2dit2dnYrdnZ2K3Z2dit2dnYrdnZ2K3Z2dit2dnYrdnZ2K3Z2dit2dnYrdnZ2K3Z2ditwcHArcHBwK3BwcCtwcHArcHBwKzMzMwUAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAe3t7VeTk5P/j4+P/4eHh/+Dg4P/f39//3t7e/93d3f/b29v/2tra/9nZ2f/Y2Nj/19fX/9XV1f/U1NT/09PT/9LS0v/R0dH/0NDQ/9DQ0P/Q0ND/xsbGpgAAAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB8fHxW7+/v/+7u7v/t7e3/7Ozs/+rq6v/p6en/2tra/87Ozv/Nzc3/zMzM/8vLy//Kysr/ysrK/8nJyf/IyMj/x8fH/8fHx//FxcX/xcXF/87Ozv/l5eX/xMTEkwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHx8fFbx8fH/8PDw/+7u7v/t7e3/7Ozs/+rq6v/o6Oj/5ubm/+Xl5f/k5OT/4+Pj/+Hh4f/h4eH/39/f/97e3v/d3d3/3Nzc/9ra2v/Z2dn/2dnZ/+fn5//o6Oj/vr6+ewAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAe3t7V/Hx8f/y8vL/8PDw/+7u7v/t7e3/7Ozs/+vr6//p6en/6Ojo/+bm5v/f39//3t7e/93d3f/c3Nz/29vb/9ra2v/e3t7/3d3d/9zc3P/b29v/6Ojo/+3t7f/m5ub+ubm5ZgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB7e3tX8fHx//Hx8f/y8vL/8PDw/+7u7v/t7e3/7Ozs/+vr6//q6ur/5+fn/9PT0//S0tL/0tLS/9HR0f/Q0ND/z8/P/97e3v/e3t7/3d3d/9zc3P/p6en/7+/v/+3t7f/j4+P8tbW1UwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHl5eVjx8fH/8fHx//Ly8v/x8fH/8PDw/+/v7//V1dX/0NDQ/8/Pz//Pz8//zs7O/83Nzf/MzMz/y8vL/8rKyv/Kysr/ycnJ/8jIyP/Hx8f/ysrK/+rq6v/w8PD/7+/v/+zs7P/i4uL3rq6uPwAAAAAAAAAAAAAAAAAAAAAAAAAAeXl5WPHx8f/x8fH/8fHx//Ly8v/x8fH/7+/v/9bW1v/R0dH/0NDQ/8/Pz//Pz8//zs7O/83Nzf/MzMz/zMzM/8vLy//Kysr/ycnJ/8jIyP/Ly8v/7Ozs/+/v7//u7u7/7e3t/+vr6//f39/wpqamMQAAAAAAAAAAAAAAAAAAAAB4eHhZ8fHx//Hx8f/x8fH/8fHx//Ly8v/w8PD/8PDw/+Hh4f/X19f/1tbW/9XV1f/U1NT/09PT/9LS0v/R0dH/0NDQ/8/Pz//Pz8//1tbW/+Dg4P/g4OD/3t7e/9zc3P/b29v/2tra/9vb2/+7u7vMAAAAAAAAAAAAAAAAAAAAAHh4eFnx8fH/8fHx//Hx8f/x8fH/8vLy//Hx8f/w8PD/5eXl/9zc3P/b29v/2tra/9nZ2f/Y2Nj/2NjY/9fX1//W1tb/1dXV/9TU1P/Z2dn/4eHh/+Dg4P/f39//3t7e/9zc3P/b29v/2tra/7a2ttQAAAAAAAAAAAAAAAAAAAAAeHh4WfDw8P/x8fH/8fHx//Hx8f/x8fH/8fHx//Dw8P/l5eX/3d3d/93d3f/c3Nz/29vb/9ra2v/Z2dn/2NjY/9fX1//W1tb/1dXV/9ra2v/j4+P/4uLi/+Dg4P/f39//3t7e/93d3f/b29v/uLi41AAAAAAAAAAAAAAAAAAAAAB4eHhZ8PDw//Dw8P/x8fH/8fHx//Hx8f/x8fH/8PDw/+Dg4P/W1tb/1tbW/9XV1f/V1dX/1NTU/9PT0//S0tL/0dHR/9DQ0P/Q0ND/2NjY/+Tk5P/j4+P/4uLi/+Dg4P/f39//3t7e/93d3f+5ubnUAAAAAAAAAAAAAAAAAAAAAHd3d1rw8PD/8PDw//Dw8P/x8fH/8fHx//Hx8f/x8fH/39/f/9LS0v/S0tL/0tLS/9LS0v/R0dH/0NDQ/9DQ0P/Pz8//zs7O/83Nzf/X19f/5ubm/+Tk5P/j4+P/4uLi/+Hh4f/f39//3t7e/7q6utQAAAAAAAAAAAAAAAAAAAAAdXV1W/Dw8P/w8PD/8PDw//Hx8f/x8fH/8fHx/9jY2P/T09P/09PT/9PT0//T09P/09PT/9PT0//S0tL/0dHR/9DQ0P/Pz8//zs7O/87Ozv/Q0ND/5ubm/+Tk5P/j4+P/4uLi/+Hh4f/g4OD/u7u71AAAAAAAAAAAAAAAAAAAAAB1dXVb8PDw//Dw8P/w8PD/8PDw//Hx8f/x8fH/4ODg/93d3f/d3d3/3Nzc/9zc3P/c3Nz/3Nzc/9zc3P/b29v/2tra/9nZ2f/Y2Nj/19fX/9nZ2f/n5+f/5ubm/+Xl5f/j4+P/4uLi/+Hh4f+8vLzUAAAAAAAAAAAAAAAAAAAAAHR0dFzw8PD/8PDw//Dw8P/w8PD/8PDw/+zs7P/e3t7/3Nzc/93d3f/d3d3/3Nzc/9zc3P/c3Nz/3Nzc/9zc3P/b29v/2tra/9nZ2f/Y2Nj/2NjY/+Tk5P/n5+f/5ubm/+Xl5f/j4+P/4uLi/76+vtQAAAAAAAAAAAAAAAAAAAAAdHR0XPDw8P/w8PD/8PDw//Dw8P/w8PD/6enp/97e3v/e3t7/3t7e/9/f3//f39//39/f/9/f3//g4OD/4ODg/+Dg4P/e3t7/3Nzc/9vb2//a2tr/4uLi/+np6f/n5+f/5ubm/+Xl5f/k5OT/vr6+1AAAAAAAAAAAAAAAAAAAAAB0dHRc8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8fHx//Hx8f/x8fH/8fHx//Ly8v/y8vL/8vLy//Ly8v/z8/P/8/Pz//Pz8//y8vL/8PDw/+3t7f/r6+v/6urq/+np6f/n5+f/5ubm/+Xl5f+/v7/UAAAAAAAAAAAAAAAAAAAAAHR0dFzw8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8fHx//Hx8f/u7u7/1NTU/9TU1P/U1NT/1NTU/9TU1P/U1NT/8PDw//Pz8//y8vL/8vLy/+7u7v/r6+v/6urq/+np6f/o6Oj/5ubm/8DAwNQAAAAAAAAAAAAAAAAAAAAAc3NzXfDw8P/w8PD/8PDw//Dw8P/w8PD/8PDw/9jY2P/V1dX/1dXV/9XV1f/V1dX/1tbW/9bW1v/W1tb/1tbW/9bW1v/W1tb/1tbW/9bW1v/b29v/8vLy//Dw8P/s7Oz/6urq/+np6f/o6Oj/wcHB1AAAAAAAAAAAAAAAAAAAAABzc3Nd8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/4ODg/93d3f/e3t7/3t7e/97e3v/e3t7/3t7e/97e3v/e3t7/39/f/9/f3//f39//39/f/+Li4v/z8/P/8vLy//Hx8f/s7Oz/6urq/+np6f/CwsLUAAAAAAAAAAAAAAAAAAAAAHFxcV7w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/l5eX/4ODg/93d3f/e3t7/3t7e/97e3v/e3t7/39/f/9/f3//f39//39/f/9/f3//i4uL/6Ojo//Pz8//z8/P/8/Pz//Hx8f/s7Oz/6urq/8TExNQAAAAAAAAAAAAAAAAAAAAAcXFxXvDw8P/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/j4+P/2tra/9ra2v/a2tr/2tra/9ra2v/b29v/29vb/9vb2//b29v/29vb/+Xl5f/z8/P/8/Pz//Pz8//z8/P/8/Pz//Hx8f/t7e3/xcXF1AAAAAAAAAAAAAAAAAAAAABwcHBf8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8PDw/9/f3//T09P/09PT/9PT0//T09P/09PT/9PT0//U1NT/1NTU/9TU1P/U1NT/4eHh//Pz8//z8/P/8/Pz//Pz8//z8/P/8/Pz//Ly8v/Hx8fUAAAAAAAAAAAAAAAAAAAAAHBwcF/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/39/f/9LS0v/S0tL/09PT/9PT0//T09P/09PT/9PT0//U1NT/1NTU/9TU1P/g4OD/8/Pz//Pz8//z8/P/8/Pz//T09P/09PT/8/Pz/8rKytQAAAAAAAAAAAAAAAAAAAAAcHBwX/Dw8P/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/i4uL/2NjY/9jY2P/Y2Nj/2dnZ/9nZ2f/Z2dn/2dnZ/9ra2v/a2tr/2tra/+Tk5P/y8vL/8/Pz//Pz8//z8/P/8/Pz//T09P/z8/P/y8vL1AAAAAAAAAAAAAAAAAAAAABwcHBf8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/5eXl/9/f3//d3d3/3d3d/93d3f/d3d3/3t7e/97e3v/e3t7/3t7e/97e3v/f39//4eHh/+bm5v/y8vL/8/Pz//Pz8//z8/P/8/Pz//T09P/Ly8vUAAAAAAAAAAAAAAAAAAAAAG9vb2Dw8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/i4uL/3d3d/93d3f/d3d3/3d3d/93d3f/d3d3/3d3d/93d3f/d3d3/3d3d/97e3v/e3t7/4+Pj//Ly8v/z8/P/8/Pz//Pz8//z8/P/8/Pz/8vLy9QAAAAAAAAAAAAAAAAAAAAAb29vYPDw8P/w8PD/8PDw//Dw8P/w8PD/8PDw/+7u7v/V1dX/1dXV/9XV1f/V1dX/1dXV/9XV1f/V1dX/1tbW/9bW1v/W1tb/1tbW/9fX1//w8PD/8vLy//Ly8v/z8/P/8/Pz//Pz8//z8/P/y8vL1AAAAAAAAAAAAAAAAAAAAABubm5h8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8PDw//Dw8P/w8PD/8fHx//Hx8f/x8fH/8fHx//Ly8v/y8vL/8vLy//Ly8v/z8/P/8/Pz//Pz8//Ly8vUAAAAAAAAAAAAAAAAAAAAAEtLS0eoqKipqKioqaqqqqiqqqqoqqqqqKqqqqiqqqqoqqqqqKurq6erq6unrKyspqysrKasrKymrq6upa6urqWurq6lrq6upa+vr6Svr6+ksbGxpLKysqOysrKjs7OzorOzs6Kzs7Ois7OzopycnIQAAAAAAAAAAAAAAAAAAAAAAAAAAgAAAAIAAAACAAAAAgAAAAIAAAACAAAAAgAAAAIAAAACAAAAAgAAAAIAAAACAAAAAgAAAAIAAAACAAAAAgAAAAIAAAACAAAAAgAAAAEAAAABAAAAAQAAAAEAAAABAAAAAQAAAAEAAAABAAAAAAAAAAAAAAAAwAAA/8AAAH/AAAB/wAAAP8AAAB/AAAAPwAAAB8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAPAAAADwAAAA8AAAAc="
# Decode the Base64 string
decoded_bytes = base64.b64decode(icon)

with tempfile.NamedTemporaryFile(delete=False, suffix=".ico") as temp_file:
    # Do whatever you need with the temporary file
    temp_file.write(decoded_bytes)
    iconpath = temp_file.name

#define main tkinter window
root = tkinter.Tk()
root.geometry('800x580')
root.title("USC Cover Page Tool")

root.iconbitmap(iconpath)
os.remove(iconpath)

font_style = ("Arial CE", 10)

style_select_label = tkinter.Label(root, text="Heading Style:", font=font_style).place(x=30, y=15)
Style_List = ["USC 1", "USC 2"]
style_var = tkinter.StringVar(root)
style_var.set(Style_List[0])
style_menu = tkinter.OptionMenu(root, style_var, *Style_List)
style_menu.place(x=118, y=10)
style_menu.config(width=12, font=('Arial CE', 10))

#Input Assignment Name
assignment_name_label = tkinter.Label(root, text="Assignment Name:", font=font_style).place(x=30, y=50)
assignment_name = tkinter.Entry(root, width=50)
assignment_name.place(x=30, y=70)

#Input Course Acronym and Number
acronymn_label = tkinter.Label(root, text="Course Acronym and Number:", font=font_style).place(x=30, y=100)
acronymn = tkinter.Entry(root, width=50)
acronymn.place(x=30, y=120)

#Input Course Name
course_name_label = tkinter.Label(root, text="Course Name:", font=font_style).place(x=30, y=150)
course_name = tkinter.Entry(root, width=50)
course_name.place(x=30, y=170)

#Input Instructor Title
instructor_title_label = tkinter.Label(root, text="Instructor Title:", font=font_style).place(x=30, y=200)
instructor_title = tkinter.Entry(root, width=50)
instructor_title.place(x=30, y=220)

#Input Student Name
student_name_label = tkinter.Label(root, text="Student Name:", font=font_style).place(x=30, y=250)
student_name = tkinter.Entry(root, width=50)
student_name.place(x=30, y=270)

def create_date(day):
    st = "ˢᵗ"
    nd = "ⁿᵈ"
    rd = "ʳᵈ"
    th = "ᵗʰ"
    if (day == 1) or (day == 21) or (day == 31):
        date_def = str(day) + st
    elif (day == 2) or (day == 22):
        date_def = str(day) + nd
    elif (day == 3) or (day == 23):
        date_def = str(day) + rd
    else:
        date_def = str(day) + th

    return date_def

# Date Selection
date_select_label = tkinter.Label(root, text="Select Date:", font=font_style).place(x=30, y=300)
Calendar = Calendar(root, pos_x=30, pos_y=320, background='gray50', arrow_thickness=1)

#Building Document Preview
preview = tkinter.Canvas(root, width=425, height=550, background="white", bd=3, relief='sunken')
preview.place(x=360, y=10)

global infinity
infinity = True
def pull_form_data(startup=True):
	while infinity:
		try:
			preview.delete("all")
			assignment = assignment_name.get()
			course_acronymn = acronymn.get()
			course_title = course_name.get()
			instructor = instructor_title.get()
			name = student_name.get()

			date = Calendar.getdate()
			dates = date.split("-")

			if dates[0] != "None":
				startup = False

			if startup:
				current_date = datetime.date.today()
				current_day = current_date.day
				date_def = create_date(int(current_day))
			else:
				current_day = dates[0]
				if current_day == "None":
					date_def = ""
				else:
					date_def = create_date(int(current_day))


			month_num = int(dates[1])
			month_def = calendar.month_name[month_num]
			year_def = dates[2]

			# set dates for global use
			global day, month, year
			day = date_def
			month = month_def
			year = year_def

			#Check Style Value
			style = style_var.get()
			if style == "USC 2":
			    preview.create_text(215, 60, text="ANDREWS UNIVERSITY AND EXTENSION PROGRAMS\n"
			                                      "        UNIVERSITY OF THE SOUTHERN CARIBBEAN\n"
			                                      "                      P.O. BOX 175, PORT OF SPAIN",
			                        font=("Times New Roman", 8))
			elif style == "USC 1":
			    preview.create_text(215, 60, text="   UNIVERSITY OF THE SOUTHERN CARIBBEAN\n"
			                                      "MARACAS ROYAL RD, ST. JOSEPH, TRINIDAD, W.I.\n"
			                                      "                   P.O. BOX 175, PORT OF SPAIN",
			                        font=("Times New Roman", 8))
			preview.create_text(215, 190, text=assignment, font=("Times New Roman", 8))
			preview.create_text(215, 280, text="               An Assignment\n"
			                                   "    Presented in Partial Fulfilment\n"
			                                   "of the Requirements for the Course",
			                    font=("Times New Roman", 8))
			preview.create_text(215, 308, text=(course_acronymn + ": " + course_title), font=("Times New Roman", 8))
			preview.create_text(215, 360, text=("INSTRUCTOR: " + instructor), font=("Times New Roman", 8))
			preview.create_text(215, 400, text="By", font=("Times New Roman", 8))
			preview.create_text(215, 420, text=name, font=("Times New Roman", 8))
			preview.create_text(215, 440, text=(date_def + " " + month_def + " " + year_def), font=("Times New Roman", 8))
			preview.create_text(345, 490, text="Approval.............................", font=("Times New Roman", 8))
			time.sleep(0.05)
			preview.update()
		except:
			pass

def create_word_doc(doc_title):
	document = Document()

	#Set Document Author
	core_properties = document.core_properties
	core_properties.author = str(student_name.get())
	core_properties.comments = "University of the Southern Caribbean"

	#Set Margins
	sections = document.sections
	for section in sections:
	    section.top_margin = Cm(2.54)
	    section.bottom_margin = Cm(2.54)
	    section.left_margin = Cm(2.54)
	    section.right_margin = Cm(2.54)

	style = document.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = Pt(12)

	style = style_var.get()
	if style == "USC 2":
	    address = "ANDREWS UNIVERSITY AND EXTENSION PROGRAMS\n" \
	              "UNIVERSITY OF THE SOUTHERN CARIBBEAN\n" \
	              "P.O. BOX 175, PORT OF SPAIN\n\n\n\n\n\n\n\n\n"
	elif style == "USC 1":
	    address = "UNIVERSITY OF THE SOUTHERN CARIBBEAN" \
	          "\n MARACAS ROYAL RD, ST. JOSEPH, TRINIDAD, W.I.\n" \
	          "P.O. BOX 175, PORT OF SPAIN\n\n\n\n\n\n\n\n\n"

	USC_Address = document.add_paragraph(address)
	paragraph_format = USC_Address.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	USC_Address.style = document.styles['Normal']

	Assignment = assignment_name.get()
	Assignment_Title = document.add_paragraph(Assignment + "\n\n\n\n\n\n\n\n\n")
	paragraph_format = Assignment_Title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Assignment_Title.style = document.styles['Normal']

	Assignment_Declaration = document.add_paragraph("An Assignment\n"
	                                            "Presented in Partial Fulfilment\n"
	                                            "of the Requirements for the Course")
	paragraph_format = Assignment_Declaration.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Assignment_Declaration.style = document.styles['Normal']

	Course_Acronym = acronymn.get()
	Course_Name = course_name.get()
	Course_Declaration = document.add_paragraph(Course_Acronym + ": " + Course_Name + "\n\n\n\n\n")
	paragraph_format = Course_Declaration.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Course_Declaration.style = document.styles['Normal']

	Instructor = instructor_title.get()
	Course_Instructor = document.add_paragraph("INSTRUCTOR: " + Instructor + "\n\n\n")
	paragraph_format = Course_Instructor.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Course_Instructor.style = document.styles['Normal']

	Student_Name = student_name.get()
	global day, month, year
	Month_Due = month
	Year_Due = year
	Date_Due = day
	Student_Declaration = document.add_paragraph("By\n\n" + Student_Name + "\n\n" + Date_Due + " " +
	                                         Month_Due + " " + Year_Due + "\n\n\n\n\n")
	paragraph_format = Student_Declaration.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Student_Declaration.style = document.styles['Normal']

	Approval_Declaration = document.add_paragraph("Approval............................")
	paragraph_format = Approval_Declaration.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	paragraph_format.space_after = 0
	paragraph_format.space_before = 0
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	Approval_Declaration.style = document.styles['Normal']

	document.save(doc_title)

def open_save_dialog():
    filename = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Microsoft Word Document", "*.docx"), ("All files", "*.*")])
    # Do something with the selected filename, like saving data to the file
    if filename:
        create_word_doc(filename)

# Create a button to trigger the save dialog
save_button = tkinter.Button(root, text="Save", width=10, command=open_save_dialog, font=font_style, background="gray80")
save_button.place(x=180, y=550, anchor='center')

#exit saftely
def safety():
    global infinity
    infinity = False
    time.sleep(0.1)
    preview.delete("all")
    preview.create_text(215, 280, text="EXITING...", font=("Times New Roman", 12))
    preview.update()
    time.sleep(1)
    root.destroy()

root.protocol('WM_DELETE_WINDOW', safety)

#call preview animation
threading.Thread(target=pull_form_data(), daemon=True).start()

# call mainloop for persistence
root.mainloop()