#USC Cover Page Generator
import tkinter
import time
import threading
import os
import shutil
from os import path
from tkinter import font
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = tkinter.Tk()
root.title("USC Cover Page Tool")
root.geometry('800x600')
root.resizable(0, 0)


Coverpg_Options = tkinter.Label(root, text="Cover Page Options").place(x=125, y=10)

#Checkboxes for Heading One or Heading Two
head = 1
def H_1():
    global head
    head = 0
    Heading_One.deselect()
    Heading_Two.select()

def H_2():
    global head
    head = 1
    Heading_Two.deselect()
    Heading_One.select()

Heading_One = tkinter.Checkbutton(root, text="Heading One", command=H_2)
Heading_One.place(x=30, y=40)
Heading_One.select()

Heading_Two = tkinter.Checkbutton(root, text="Heading Two", command=H_1)
Heading_Two.place(x=230, y=40)


#Input Assignment Name
assignment_name_label = tkinter.Label(root, text="Assignment Name:").place(x=30, y=80)
assignment_name = tkinter.Entry(root, width=50)
assignment_name.place(x=30, y=100)

#Input Course Acronym and Number
acronymn_label = tkinter.Label(root, text="Course Acronym and Number:").place(x=30, y=130)
acronymn = tkinter.Entry(root, width=30)
acronymn.place(x=30, y=150)

#Input Course Name
course_name_label = tkinter.Label(root, text="Course Name:").place(x=30, y=180)
course_name = tkinter.Entry(root, width=50)
course_name.place(x=30, y=200)

#Input Instructor Title
instructor_title_label = tkinter.Label(root, text="Instructor Title:").place(x=30, y=230)
instructor_title = tkinter.Entry(root, width=50)
instructor_title.place(x=30, y=250)

#Input Student Name
student_name_label = tkinter.Label(root, text="Student Name:").place(x=30, y=280)
student_name = tkinter.Entry(root, width=50)
student_name.place(x=30, y=300)

#Input Date
date_def = "1ˢᵗ"
def create_date():
    global date_def
    spinbox_val = int(date.get())
    st = "ˢᵗ"
    nd = "ⁿᵈ"
    rd = "ʳᵈ"
    th = "ᵗʰ"
    if (spinbox_val == 1) or (spinbox_val == 21) or (spinbox_val == 31):
        date_def = str(spinbox_val) + st
    elif (spinbox_val == 2) or (spinbox_val == 22):
        date_def = str(spinbox_val) + nd
    elif (spinbox_val == 3) or (spinbox_val == 23):
        date_def = str(spinbox_val) + rd
    else:
        date_def = str(spinbox_val) + th

date_label = tkinter.Label(root, text="Enter Day:").place(x=30, y=330)
date = tkinter.Spinbox(root, width=3, from_=1, to=31, state="readonly", command=create_date)
date.place(x=91, y=330)

month_label = tkinter.Label(root, text="Enter Month:").place(x=30, y=360)
month = tkinter.Entry(root, width=20)
month.place(x=105, y=360)

year_label = tkinter.Label(root, text="Enter Year:").place(x=30, y=390)
year = tkinter.Entry(root, width=20)
year.place(x=92, y=390)


#Building Document Preview
preview_label = tkinter.Label(root, text="Document Preview").place(x=525, y=10)
preview = tkinter.Canvas(root, width=425, height=550, background="white", bd=3, relief='sunken')
preview.place(x=360, y=30)

infinity = 0
def pull_form_data():
    while infinity == 0:
        preview.delete("all")
        assignment = assignment_name.get()
        course_acronymn = acronymn.get()
        course_title = course_name.get()
        instructor = instructor_title.get()
        name = student_name.get()
        month_def = month.get()
        year_def = year.get()
        #Check Heading Value
        if head == 1:
            preview.create_text(215, 60, text="ANDREWS UNIVERSITY AND EXTENSION PROGRAMS\n"
                                              "        UNIVERSITY OF THE SOUTHERN CARIBBEAN\n"
                                              "                          P.O. BOX 175, PORT OF SPAIN",
                                font=("Times New Roman", 6))
        else:
            preview.create_text(215, 60, text="   UNIVERSITY OF THE SOUTHERN CARIBBEAN\n"
                                              "MARACAS ROYAL ROAD, MARACUS, ST. JOSEPH\n"
                                              "                   P.O. BOX 175, PORT OF SPAIN",
                                font=("Times New Roman", 6))
        preview.create_text(215, 190, text=assignment, font=("Times New Roman", 6))
        preview.create_text(215, 280, text="               An Assignment\n"
                                           "    Presented in Partial Fulfilment\n"
                                           "Of the Requirements of the Course",
                            font=("Times New Roman", 6))
        preview.create_text(215, 300, text=(course_acronymn + ": " + course_title), font=("Times New Roman", 6))
        preview.create_text(215, 360, text=("INSTRUCTOR: " + instructor), font=("Times New Roman", 6))
        preview.create_text(215, 400, text="By", font=("Times New Roman", 6))
        preview.create_text(215, 420, text=name, font=("Times New Roman", 6))
        preview.create_text(215, 440, text=(date_def + " " + month_def + " " + year_def), font=("Times New Roman", 6))
        preview.create_text(345, 490, text="Approval.............................", font=("Times New Roman", 6))
        time.sleep(0.05)
        preview.update()


def cover_page():
    global infinity
    infinity = 1
    time.sleep(0.1)
    preview.delete("all")
    preview.create_text(215, 280, text="SAVING...", font=("Times New Roman", 12))
    preview.update()
    document = Document()

    #Set Margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    if head == 1:
        address = "ANDREWS UNIVERSITY AND EXTENSION PROGRAMS\n" \
                  "UNIVERSITY OF THE SOUTHERN CARIBBEAN\n" \
                  "P.O. BOX 175, PORT OF SPAIN\n\n\n\n\n\n\n\n\n"
    else:
        address = "UNIVERSITY OF THE SOUTHERN CARIBBEAN" \
              "\n MARACAS ROYAL ROAD, MARACUS, ST. JOSEPH\n" \
              "P.O. BOX 175, PORT OF SPAIN\n\n\n\n\n\n\n\n\n"
    USC_Address = document.add_paragraph(address)
    paragraph_format = USC_Address.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    USC_Address.style = document.styles['Normal']

    Assignment = assignment_name.get()
    Assignment_Title = document.add_paragraph(Assignment + "\n\n\n\n\n\n\n\n\n")
    paragraph_format = Assignment_Title.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Assignment_Title.style = document.styles['Normal']

    Assignment_Declaration = document.add_paragraph("An Assignment\n"
                                                "Presented in Partial Fulfilment\n"
                                                "Of the Requirements of the Course")
    paragraph_format = Assignment_Declaration.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Assignment_Declaration.style = document.styles['Normal']

    Course_Acronym = acronymn.get()
    Course_Name = course_name.get()
    Course_Declaration = document.add_paragraph(Course_Acronym + ": " + Course_Name + "\n\n\n\n\n")
    paragraph_format = Course_Declaration.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Course_Declaration.style = document.styles['Normal']

    Instructor = instructor_title.get()
    Course_Instructor = document.add_paragraph("INSTRUCTOR: " + Instructor + "\n\n\n")
    paragraph_format = Course_Instructor.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Course_Instructor.style = document.styles['Normal']

    Student_Name = student_name.get()
    Month_Due = month.get()
    Year_Due = year.get()
    Date_Due = date_def
    Student_Declaration = document.add_paragraph("By\n\n" + Student_Name + "\n\n" + Date_Due + " " +
                                             Month_Due + " " + Year_Due + "\n\n\n\n\n")
    paragraph_format = Student_Declaration.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Student_Declaration.style = document.styles['Normal']

    Approval_Declaration = document.add_paragraph("Approval............................")
    paragraph_format = Approval_Declaration.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    Approval_Declaration.style = document.styles['Normal']
    doc_title = savename.get()
    if doc_title == "":
        doc_title = "USC Cover Page.docx"
    else:
        doc_title = doc_title + ".docx"
    document.save(doc_title)
    cur_dir = os.getcwd()
    cur_location = cur_dir + "/" + doc_title
    if path.exists(cur_location):
        location = savar.get()
        homepath = (os.environ['HOMEPATH'])
        if location == "Desktop":
            Desktop = homepath + ("/Desktop") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Desktop")):
                shutil.move(cur_location, Desktop)
            else:
                preview.create_text(215, 320, text="Could not locate Desktop\nFile stored in current directory")
                preview.update()
        elif location == "Downloads":
            Downloads = homepath + ("/Downloads") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Downloads")):
                shutil.move(cur_location, Downloads)
            else:
                preview.create_text(215, 320, text="Could not locate Downloads\nFile stored in current directory")
                preview.update()
        elif location == "Videos":
            Videos = homepath + ("/Videos") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Videos")):
                shutil.move(cur_location, Videos)
            else:
                preview.create_text(215, 320, text="Could not locate Videos\nFile stored in current directory")
                preview.update()
        elif location == "Music":
            Music = homepath + ("/Music") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Music")):
                shutil.move(cur_location, Music)
            else:
                preview.create_text(215, 320, text="Could not locate Music\nFile stored in current directory")
                preview.update()
        elif location == "Documents":
            Documents = homepath + ("/Documents") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Documents")):
                shutil.move(cur_location, Documents)
            else:
                preview.create_text(215, 320, text="Could not locate Documents\nFile stored in current directory")
                preview.update()
        elif location == "Pictures":
            Pictures = homepath + ("/Pictures") + "/" + doc_title + ".docx"
            if path.exists(homepath + ("/Pictures")):
                shutil.move(cur_location, Pictures)
            else:
                preview.create_text(215, 320, text="Could not locate Pictures\nFile stored in current directory")
                preview.update()
    time.sleep(2)
    infinity = 0


#Save Button
simple_canvas = tkinter.Canvas(root, width=300, height=120, background="lightgrey")
simple_canvas.place(x=30, y=460)
simple_canvas.create_text(60, 20, text="Save As:")
simple_canvas.create_text(60, 55, text="Location:")
savename = tkinter.Entry(root, width=20)
savename.place(x=115, y=470)
savename.insert(0, 'USC Cover Page')
Save_Document = tkinter.Label(root, text="Save Settings").place(x=140, y=440)
save = tkinter.Button(root, text="Save", command=cover_page)
save_font = font.Font(size=12)
save['font'] = save_font
save.place(x=150, y=540)

homepath = (os.environ['HOMEPATH'])
if path.exists(homepath):
    Documents = homepath + ("/Documents")
    Pictures = homepath + ("/Pictures")
    SaveList = ["Desktop", "Documents", "Pictures", "Music", "Videos", "Downloads"]
else:
    SaveList = ["Current Directory"]
savar = tkinter.StringVar(root)
savar.set(SaveList[0])
save_menu = tkinter.OptionMenu(root, savar, *SaveList)
save_menu.place(x=118, y=500)
save_menu.config(width=12, font=('Times New Roman', 8))

#exit saftely
def safety():
    global infinity
    infinity = 1
    time.sleep(0.1)
    preview.delete("all")
    preview.create_text(215, 280, text="EXITING...", font=("Times New Roman", 12))
    preview.update()
    time.sleep(1)
    root.destroy()

root.protocol('WM_DELETE_WINDOW', safety)

#call preview animation
threading.Thread(target=pull_form_data(), daemon=True).start()

root.mainloop()